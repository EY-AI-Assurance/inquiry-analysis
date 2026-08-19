from __future__ import annotations

from io import BytesIO

from docx import Document

from .common import (
    CorruptDocumentError,
    EmptyDocumentError,
    ParsedDocument,
    SourceChunk,
    column_name,
    normalize_text,
)


def parse_docx(data: bytes) -> ParsedDocument:
    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise CorruptDocumentError("Word 文件无法读取，可能已损坏或格式不正确。") from exc

    chunks: list[SourceChunk] = []
    paragraph_buffer: list[str] = []
    paragraph_start = 1
    order = 0

    def flush_paragraphs(end_index: int) -> None:
        nonlocal paragraph_buffer, paragraph_start, order
        if not paragraph_buffer:
            paragraph_start = end_index + 1
            return
        order += 1
        chunks.append(
            SourceChunk(
                source_id=f"word-paragraphs-{paragraph_start}-{end_index}",
                locator=f"Word 段落 {paragraph_start}–{end_index}",
                content="\n".join(paragraph_buffer),
                order=order,
            )
        )
        paragraph_buffer = []
        paragraph_start = end_index + 1

    for index, paragraph in enumerate(document.paragraphs, start=1):
        value = normalize_text(paragraph.text)
        if value:
            paragraph_buffer.append(f"[P{index}] {value}")
        if len(paragraph_buffer) >= 20 or sum(map(len, paragraph_buffer)) >= 4_000:
            flush_paragraphs(index)
    flush_paragraphs(len(document.paragraphs))

    for table_index, table in enumerate(document.tables, start=1):
        lines = []
        for row_index, row in enumerate(table.rows, start=1):
            cells = [normalize_text(cell.text).replace("\n", " ") for cell in row.cells]
            if any(cells):
                fields = [
                    f"{column_name(column_index)}（T{table_index}:R{row_index}:C{column_index}）={value}"
                    for column_index, value in enumerate(cells, start=1)
                    if value
                ]
                lines.append(f"第 {row_index} 行 | " + " | ".join(fields))
        if lines:
            order += 1
            chunks.append(
                SourceChunk(
                    source_id=f"word-table-{table_index}",
                    locator=f"Word 表格 {table_index}",
                    content="\n".join(lines),
                    order=order,
                )
            )

    if not chunks:
        raise EmptyDocumentError("Word 文件中没有可分析的文本或表格。")
    return ParsedDocument(chunks=chunks, warnings=[])
