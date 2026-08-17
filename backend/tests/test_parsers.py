from __future__ import annotations

from io import BytesIO

import openpyxl
import pytest
import xlwt
from docx import Document
from reportlab.pdfgen import canvas

from parsers import ScannedPdfError, UnsupportedFileError, parse_document


def test_csv_parser_preserves_row_locations():
    parsed = parse_document(
        "profit-and-loss.csv",
        "项目,2025,2024\n收入,1200,1000\n经营利润,80,45\n".encode(),
    )
    assert parsed.chunks[0].source_id == "csv-rows-1-3"
    assert parsed.chunks[0].locator == "CSV 行 1–3"
    assert "项目（A2）=收入" in parsed.chunks[0].content
    assert "2025（B3）=80" in parsed.chunks[0].content


def test_docx_parser_reads_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("2025 年收入增长 20%")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "金额"
    table.cell(1, 0).text = "Adjusted EBITDA"
    table.cell(1, 1).text = "100"
    stream = BytesIO()
    document.save(stream)

    parsed = parse_document("report.docx", stream.getvalue())
    assert len(parsed.chunks) == 2
    assert parsed.chunks[1].locator == "Word 表格 1"


def test_xlsx_parser_reads_sheet_cells():
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "损益表"
    sheet.append(["项目", "本期"])
    sheet.append(["净利润", -10])
    stream = BytesIO()
    workbook.save(stream)

    parsed = parse_document("report.xlsx", stream.getvalue())
    assert parsed.chunks[0].locator == "工作表“损益表”行 1–2"
    assert "本期（B2）=-10" in parsed.chunks[0].content


def test_legacy_xls_parser():
    workbook = xlwt.Workbook()
    sheet = workbook.add_sheet("P&L")
    sheet.write(0, 0, "Revenue")
    sheet.write(1, 0, 100)
    stream = BytesIO()
    workbook.save(stream)

    parsed = parse_document("report.xls", stream.getvalue())
    assert parsed.chunks[0].locator == "工作表“P&L”行 1–2"


def test_pdf_parser_tracks_page_number():
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.drawString(72, 720, "Revenue 2025: 1,200")
    pdf.showPage()
    pdf.save()

    parsed = parse_document("report.pdf", stream.getvalue())
    assert parsed.chunks[0].locator == "PDF 第 1 页"
    assert "Revenue" in parsed.chunks[0].content


def test_scanned_or_blank_pdf_is_rejected():
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.showPage()
    pdf.save()

    with pytest.raises(ScannedPdfError, match="OCR"):
        parse_document("scan.pdf", stream.getvalue())


def test_old_doc_has_actionable_message():
    with pytest.raises(UnsupportedFileError, match=".docx"):
        parse_document("old-report.doc", b"old word data")


def test_large_document_is_trimmed_with_warning():
    rows = ["项目,金额"]
    rows.extend(f"收入项目{i},{i}" for i in range(5_000))
    parsed = parse_document("large.csv", "\n".join(rows).encode(), max_chars=8_000)
    assert sum(len(chunk.content) for chunk in parsed.chunks) <= 8_000
    assert parsed.warnings
